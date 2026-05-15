#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Генерація DOCX зі схемою БД DevFlow (джерело: packages/backend/src/scripts/migrate.js)."""

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

OUT = Path(__file__).resolve().parent.parent / "DevFlow-shema-bazy.docx"

# Колонка: (ім'я, тип, обмеження/примітка)
# FK: (колонка, ref_table, ref_column, on_delete)
# IDX: рядок опису індексу

TABLES = [
    {
        "name": "users",
        "group": "Користувачі",
        "desc": "Облікові записи (email/пароль або GitHub OAuth).",
        "columns": [
            ("id", "INT", "PK, AUTO_INCREMENT"),
            ("username", "VARCHAR(30)", "UNIQUE, NOT NULL"),
            ("email", "VARCHAR(255)", "UNIQUE, NOT NULL"),
            ("password", "VARCHAR(255)", "NULL (GitHub-only без пароля)"),
            ("reputation", "INT", "DEFAULT 0"),
            ("role", "ENUM / VARCHAR(20)", "user | moderator | admin"),
            ("bio", "TEXT", ""),
            ("location", "VARCHAR(100)", ""),
            ("website", "VARCHAR(255)", ""),
            ("avatar_url", "VARCHAR(500)", "NULL"),
            ("github_id", "BIGINT", "UNIQUE, NULL"),
            ("github_login", "VARCHAR(64)", "NULL"),
            ("github_avatar_url", "VARCHAR(500)", "NULL"),
            ("github_access_token", "TEXT", "NULL"),
            ("github_profile", "JSON", "NULL"),
            ("github_stack", "JSON", "NULL"),
            ("github_contributions", "JSON", "NULL"),
            ("github_badges", "JSON", "NULL"),
            ("github_synced_at", "DATETIME", "NULL"),
            ("created_at", "DATETIME", "NOT NULL"),
            ("updated_at", "DATETIME", "NOT NULL"),
        ],
        "indexes": ["idx_username", "idx_email", "idx_reputation", "idx_github_login", "uq_github_id"],
        "fks": [],
    },
    {
        "name": "user_repositories",
        "group": "Користувачі",
        "desc": "Кеш репозиторіїв GitHub користувача.",
        "columns": [
            ("id", "INT", "PK"),
            ("user_id", "INT", "NOT NULL → users.id"),
            ("github_repo_id", "BIGINT", "NOT NULL"),
            ("name", "VARCHAR(120)", "NOT NULL"),
            ("full_name", "VARCHAR(200)", "NOT NULL"),
            ("html_url", "VARCHAR(500)", "NOT NULL"),
            ("description", "TEXT", "NULL"),
            ("homepage", "VARCHAR(500)", "NULL"),
            ("language", "VARCHAR(60)", "NULL"),
            ("languages", "JSON", "NULL"),
            ("topics", "JSON", "NULL"),
            ("stars", "INT", "DEFAULT 0"),
            ("forks", "INT", "DEFAULT 0"),
            ("watchers", "INT", "DEFAULT 0"),
            ("open_issues", "INT", "DEFAULT 0"),
            ("is_fork", "BOOLEAN", "DEFAULT FALSE"),
            ("is_archived", "BOOLEAN", "DEFAULT FALSE"),
            ("is_pinned", "BOOLEAN", "DEFAULT FALSE"),
            ("pushed_at", "DATETIME", "NULL"),
            ("repo_created_at", "DATETIME", "NULL"),
            ("synced_at", "DATETIME", "NOT NULL"),
        ],
        "indexes": ["uq_user_repo (user_id, github_repo_id)", "idx_user", "idx_pinned", "idx_lang"],
        "fks": [("user_id", "users", "id", "CASCADE")],
    },
    {
        "name": "questions",
        "group": "Legacy Q&A",
        "desc": "Старі питання (окрема від content_items).",
        "columns": [
            ("id", "INT", "PK"),
            ("title", "VARCHAR(255)", "NOT NULL"),
            ("body", "TEXT", "NOT NULL"),
            ("tags", "JSON", "NOT NULL"),
            ("author_id", "INT", "NOT NULL → users.id"),
            ("views", "INT", "DEFAULT 0"),
            ("upvotes", "INT", "DEFAULT 0"),
            ("downvotes", "INT", "DEFAULT 0"),
            ("created_at", "DATETIME", "NOT NULL"),
            ("updated_at", "DATETIME", "NOT NULL"),
        ],
        "indexes": ["idx_author", "idx_created", "idx_views", "FULLTEXT(title, body)"],
        "fks": [("author_id", "users", "id", "CASCADE")],
    },
    {
        "name": "answers",
        "group": "Legacy Q&A",
        "desc": "Відповіді до legacy questions.",
        "columns": [
            ("id", "INT", "PK"),
            ("body", "TEXT", "NOT NULL"),
            ("question_id", "INT", "NOT NULL → questions.id"),
            ("author_id", "INT", "NOT NULL → users.id"),
            ("is_accepted", "BOOLEAN", "DEFAULT FALSE"),
            ("upvotes", "INT", "DEFAULT 0"),
            ("downvotes", "INT", "DEFAULT 0"),
            ("created_at", "DATETIME", "NOT NULL"),
            ("updated_at", "DATETIME", "NOT NULL"),
        ],
        "indexes": ["idx_question", "idx_author", "idx_accepted", "FULLTEXT(body)"],
        "fks": [
            ("question_id", "questions", "id", "CASCADE"),
            ("author_id", "users", "id", "CASCADE"),
        ],
    },
    {
        "name": "votes",
        "group": "Спільне",
        "desc": "Голоси up/down для question, answer, content, content_answer.",
        "columns": [
            ("id", "INT", "PK"),
            ("user_id", "INT", "NOT NULL → users.id"),
            ("entity_type", "ENUM", "question | answer | content | content_answer"),
            ("entity_id", "INT", "NOT NULL (поліморфний ID)"),
            ("vote_type", "ENUM", "up | down"),
            ("created_at", "DATETIME", "NOT NULL"),
        ],
        "indexes": ["unique_vote (user_id, entity_type, entity_id)", "idx_entity"],
        "fks": [("user_id", "users", "id", "CASCADE")],
    },
    {
        "name": "notifications",
        "group": "Спільне",
        "desc": "Сповіщення користувача (type/entity_type — VARCHAR після міграції).",
        "columns": [
            ("id", "INT", "PK"),
            ("user_id", "INT", "NOT NULL → users.id"),
            ("type", "VARCHAR(64)", "NOT NULL"),
            ("entity_type", "VARCHAR(64)", "NOT NULL"),
            ("entity_id", "INT", "NOT NULL"),
            ("data", "JSON", "NULL"),
            ("is_read", "BOOLEAN", "DEFAULT FALSE"),
            ("actor_id", "INT", "NULL (без FK у migrate)"),
            ("created_at", "DATETIME", "NOT NULL"),
        ],
        "indexes": ["idx_user", "idx_read", "idx_created"],
        "fks": [("user_id", "users", "id", "CASCADE")],
    },
    {
        "name": "bookmarks",
        "group": "Legacy Q&A",
        "desc": "Закладки на legacy questions.",
        "columns": [
            ("id", "INT", "PK"),
            ("user_id", "INT", "NOT NULL"),
            ("question_id", "INT", "NOT NULL"),
            ("created_at", "DATETIME", "NOT NULL"),
        ],
        "indexes": ["unique_bookmark (user_id, question_id)", "idx_user", "idx_question"],
        "fks": [
            ("user_id", "users", "id", "CASCADE"),
            ("question_id", "questions", "id", "CASCADE"),
        ],
    },
    {
        "name": "question_views",
        "group": "Legacy Q&A",
        "desc": "Унікальні перегляди питання (viewer_key = user або visitor).",
        "columns": [
            ("id", "INT", "PK"),
            ("question_id", "INT", "NOT NULL"),
            ("viewer_key", "VARCHAR(96)", "NOT NULL"),
            ("viewed_at", "DATETIME", "NOT NULL"),
        ],
        "indexes": ["uq_question_viewer (question_id, viewer_key)", "idx_question"],
        "fks": [("question_id", "questions", "id", "CASCADE")],
    },
    {
        "name": "articles",
        "group": "Knowledge Hub (окремі таблиці)",
        "desc": "Статті.",
        "columns": [
            ("id", "INT", "PK"),
            ("title", "VARCHAR(255)", "NOT NULL"),
            ("body", "TEXT", "NOT NULL"),
            ("excerpt", "VARCHAR(280)", "NOT NULL"),
            ("tags", "JSON", "NOT NULL"),
            ("author_id", "INT", "NOT NULL"),
            ("views", "INT", "DEFAULT 0"),
            ("created_at", "DATETIME", "NOT NULL"),
            ("updated_at", "DATETIME", "NOT NULL"),
        ],
        "indexes": ["idx_author", "idx_created", "idx_views", "FULLTEXT(title, body, excerpt)"],
        "fks": [("author_id", "users", "id", "CASCADE")],
    },
    {
        "name": "article_views",
        "group": "Knowledge Hub (окремі таблиці)",
        "desc": "Унікальні перегляди статей.",
        "columns": [
            ("id", "INT", "PK"),
            ("article_id", "INT", "NOT NULL"),
            ("viewer_key", "VARCHAR(96)", "NOT NULL"),
            ("viewed_at", "DATETIME", "NOT NULL"),
        ],
        "indexes": ["uq_article_viewer", "idx_article"],
        "fks": [("article_id", "articles", "id", "CASCADE")],
    },
    {
        "name": "guides",
        "group": "Knowledge Hub (окремі таблиці)",
        "desc": "Гайди.",
        "columns": [
            ("id", "INT", "PK"),
            ("title", "VARCHAR(255)", "NOT NULL"),
            ("summary", "VARCHAR(280)", "NOT NULL"),
            ("excerpt", "VARCHAR(280)", "NOT NULL"),
            ("body", "TEXT", "NOT NULL"),
            ("difficulty", "ENUM", "beginner | intermediate | advanced"),
            ("estimated_minutes", "INT", "DEFAULT 15"),
            ("tags", "JSON", "NOT NULL"),
            ("author_id", "INT", "NOT NULL"),
            ("views", "INT", "DEFAULT 0"),
            ("created_at", "DATETIME", "NOT NULL"),
            ("updated_at", "DATETIME", "NOT NULL"),
        ],
        "indexes": ["idx_author", "idx_created", "idx_views", "idx_difficulty", "FULLTEXT"],
        "fks": [("author_id", "users", "id", "CASCADE")],
    },
    {
        "name": "guide_views",
        "group": "Knowledge Hub (окремі таблиці)",
        "columns": [
            ("id", "INT", "PK"),
            ("guide_id", "INT", "NOT NULL"),
            ("viewer_key", "VARCHAR(96)", "NOT NULL"),
            ("viewed_at", "DATETIME", "NOT NULL"),
        ],
        "indexes": ["uq_guide_viewer", "idx_guide"],
        "fks": [("guide_id", "guides", "id", "CASCADE")],
    },
    {
        "name": "snippets",
        "group": "Knowledge Hub (окремі таблиці)",
        "desc": "Сніпети коду.",
        "columns": [
            ("id", "INT", "PK"),
            ("title", "VARCHAR(255)", "NOT NULL"),
            ("description", "TEXT", "NOT NULL"),
            ("excerpt", "VARCHAR(280)", "NOT NULL"),
            ("code", "MEDIUMTEXT", "NOT NULL"),
            ("language", "VARCHAR(40)", "NOT NULL"),
            ("tags", "JSON", "NOT NULL"),
            ("author_id", "INT", "NOT NULL"),
            ("views", "INT", "DEFAULT 0"),
            ("created_at", "DATETIME", "NOT NULL"),
            ("updated_at", "DATETIME", "NOT NULL"),
        ],
        "indexes": ["idx_author", "idx_language", "FULLTEXT"],
        "fks": [("author_id", "users", "id", "CASCADE")],
    },
    {
        "name": "snippet_views",
        "group": "Knowledge Hub (окремі таблиці)",
        "columns": [
            ("id", "INT", "PK"),
            ("snippet_id", "INT", "NOT NULL"),
            ("viewer_key", "VARCHAR(96)", "NOT NULL"),
            ("viewed_at", "DATETIME", "NOT NULL"),
        ],
        "indexes": ["uq_snippet_viewer", "idx_snippet"],
        "fks": [("snippet_id", "snippets", "id", "CASCADE")],
    },
    {
        "name": "roadmaps",
        "group": "Knowledge Hub (окремі таблиці)",
        "columns": [
            ("id", "INT", "PK"),
            ("title", "VARCHAR(255)", "NOT NULL"),
            ("summary", "VARCHAR(280)", "NOT NULL"),
            ("excerpt", "VARCHAR(280)", "NOT NULL"),
            ("body", "TEXT", "NOT NULL"),
            ("steps", "JSON", "NOT NULL"),
            ("difficulty", "ENUM", "beginner | intermediate | advanced"),
            ("estimated_weeks", "INT", "DEFAULT 4"),
            ("tags", "JSON", "NOT NULL"),
            ("author_id", "INT", "NOT NULL"),
            ("views", "INT", "DEFAULT 0"),
            ("created_at", "DATETIME", "NOT NULL"),
            ("updated_at", "DATETIME", "NOT NULL"),
        ],
        "indexes": ["idx_author", "idx_difficulty", "FULLTEXT"],
        "fks": [("author_id", "users", "id", "CASCADE")],
    },
    {
        "name": "roadmap_views",
        "group": "Knowledge Hub (окремі таблиці)",
        "columns": [
            ("id", "INT", "PK"),
            ("roadmap_id", "INT", "NOT NULL"),
            ("viewer_key", "VARCHAR(96)", "NOT NULL"),
            ("viewed_at", "DATETIME", "NOT NULL"),
        ],
        "indexes": ["uq_roadmap_viewer", "idx_roadmap"],
        "fks": [("roadmap_id", "roadmaps", "id", "CASCADE")],
    },
    {
        "name": "best_practices",
        "group": "Knowledge Hub (окремі таблиці)",
        "columns": [
            ("id", "INT", "PK"),
            ("title", "VARCHAR(255)", "NOT NULL"),
            ("rule", "VARCHAR(500)", "NOT NULL"),
            ("excerpt", "VARCHAR(280)", "NOT NULL"),
            ("body", "TEXT", "NOT NULL"),
            ("anti_patterns", "TEXT", "NULL"),
            ("category", "VARCHAR(80)", "NULL"),
            ("tags", "JSON", "NOT NULL"),
            ("author_id", "INT", "NOT NULL"),
            ("views", "INT", "DEFAULT 0"),
            ("created_at", "DATETIME", "NOT NULL"),
            ("updated_at", "DATETIME", "NOT NULL"),
        ],
        "indexes": ["idx_category", "FULLTEXT"],
        "fks": [("author_id", "users", "id", "CASCADE")],
    },
    {
        "name": "best_practice_views",
        "group": "Knowledge Hub (окремі таблиці)",
        "columns": [
            ("id", "INT", "PK"),
            ("best_practice_id", "INT", "NOT NULL"),
            ("viewer_key", "VARCHAR(96)", "NOT NULL"),
            ("viewed_at", "DATETIME", "NOT NULL"),
        ],
        "indexes": ["uq_bp_viewer", "idx_bp"],
        "fks": [("best_practice_id", "best_practices", "id", "CASCADE")],
    },
    {
        "name": "faqs",
        "group": "Knowledge Hub (окремі таблиці)",
        "columns": [
            ("id", "INT", "PK"),
            ("title", "VARCHAR(255)", "NOT NULL"),
            ("topic", "VARCHAR(120)", "NOT NULL"),
            ("excerpt", "VARCHAR(280)", "NOT NULL"),
            ("body", "TEXT", "NOT NULL"),
            ("qa_pairs", "JSON", "NOT NULL"),
            ("tags", "JSON", "NOT NULL"),
            ("author_id", "INT", "NOT NULL"),
            ("views", "INT", "DEFAULT 0"),
            ("created_at", "DATETIME", "NOT NULL"),
            ("updated_at", "DATETIME", "NOT NULL"),
        ],
        "indexes": ["idx_topic", "FULLTEXT"],
        "fks": [("author_id", "users", "id", "CASCADE")],
    },
    {
        "name": "faq_views",
        "group": "Knowledge Hub (окремі таблиці)",
        "columns": [
            ("id", "INT", "PK"),
            ("faq_id", "INT", "NOT NULL"),
            ("viewer_key", "VARCHAR(96)", "NOT NULL"),
            ("viewed_at", "DATETIME", "NOT NULL"),
        ],
        "indexes": ["uq_faq_viewer", "idx_faq"],
        "fks": [("faq_id", "faqs", "id", "CASCADE")],
    },
    {
        "name": "content_items",
        "group": "Knowledge Hub (ядро)",
        "desc": "Універсальне сховище hub-контенту (питання, статті, гайди тощо).",
        "columns": [
            ("id", "INT", "PK"),
            ("type", "ENUM", "question | article | guide | snippet | roadmap | best_practice | faq"),
            ("title", "VARCHAR(255)", "NOT NULL"),
            ("slug", "VARCHAR(280)", "UNIQUE, NOT NULL"),
            ("body", "MEDIUMTEXT", "NOT NULL"),
            ("excerpt", "VARCHAR(500)", "NULL"),
            ("tags", "JSON", "NULL"),
            ("author_id", "INT", "NOT NULL"),
            ("status", "ENUM", "draft | published | archived"),
            ("difficulty", "ENUM", "beginner | intermediate | advanced, NULL"),
            ("technology", "VARCHAR(80)", "NULL"),
            ("estimated_read_time", "INT", "NULL"),
            ("meta", "JSON", "NULL"),
            ("views", "INT", "DEFAULT 0"),
            ("is_featured", "BOOLEAN", "DEFAULT FALSE"),
            ("created_at", "DATETIME", "NOT NULL"),
            ("updated_at", "DATETIME", "NOT NULL"),
            ("published_at", "DATETIME", "NULL"),
        ],
        "indexes": ["uq_slug", "idx_type", "idx_type_published", "idx_author", "FULLTEXT"],
        "fks": [("author_id", "users", "id", "CASCADE")],
    },
    {
        "name": "content_answers",
        "group": "Knowledge Hub (ядро)",
        "desc": "Відповіді на content_items типу question.",
        "columns": [
            ("id", "INT", "PK"),
            ("content_id", "INT", "NOT NULL → content_items.id"),
            ("body", "MEDIUMTEXT", "NOT NULL"),
            ("author_id", "INT", "NOT NULL"),
            ("is_accepted", "BOOLEAN", "DEFAULT FALSE"),
            ("created_at", "DATETIME", "NOT NULL"),
            ("updated_at", "DATETIME", "NOT NULL"),
        ],
        "indexes": ["idx_content", "idx_author", "idx_accepted", "FULLTEXT(body)"],
        "fks": [
            ("content_id", "content_items", "id", "CASCADE"),
            ("author_id", "users", "id", "CASCADE"),
        ],
    },
    {
        "name": "content_views",
        "group": "Knowledge Hub (ядро)",
        "columns": [
            ("id", "INT", "PK"),
            ("content_id", "INT", "NOT NULL"),
            ("viewer_key", "VARCHAR(96)", "NOT NULL"),
            ("viewed_at", "DATETIME", "NOT NULL"),
        ],
        "indexes": ["uq_content_viewer", "idx_content"],
        "fks": [("content_id", "content_items", "id", "CASCADE")],
    },
    {
        "name": "content_bookmarks",
        "group": "Knowledge Hub (ядро)",
        "columns": [
            ("id", "INT", "PK"),
            ("user_id", "INT", "NOT NULL"),
            ("content_id", "INT", "NOT NULL"),
            ("created_at", "DATETIME", "NOT NULL"),
        ],
        "indexes": ["unique_bookmark (user_id, content_id)"],
        "fks": [
            ("user_id", "users", "id", "CASCADE"),
            ("content_id", "content_items", "id", "CASCADE"),
        ],
    },
    {
        "name": "content_linked_repos",
        "group": "Knowledge Hub (ядро)",
        "desc": "GitHub-репозиторії, прив’язані до контенту (поліморфно: target_type + target_id).",
        "columns": [
            ("id", "INT", "PK"),
            ("target_type", "ENUM", "question | article | guide | snippet | roadmap | best_practice | faq | content"),
            ("target_id", "INT", "NOT NULL (без FK)"),
            ("github_repo_id", "BIGINT", "NOT NULL"),
            ("name", "VARCHAR(120)", "NOT NULL"),
            ("full_name", "VARCHAR(200)", "NOT NULL"),
            ("html_url", "VARCHAR(500)", "NOT NULL"),
            ("description", "TEXT", "NULL"),
            ("homepage", "VARCHAR(500)", "NULL"),
            ("language", "VARCHAR(60)", "NULL"),
            ("topics", "JSON", "NULL"),
            ("stars", "INT", "DEFAULT 0"),
            ("forks", "INT", "DEFAULT 0"),
            ("open_issues", "INT", "DEFAULT 0"),
            ("is_fork", "BOOLEAN", "DEFAULT FALSE"),
            ("is_archived", "BOOLEAN", "DEFAULT FALSE"),
            ("added_by_user_id", "INT", "NOT NULL"),
            ("added_note", "VARCHAR(280)", "NULL"),
            ("pushed_at", "DATETIME", "NULL"),
            ("added_at", "DATETIME", "NOT NULL"),
        ],
        "indexes": ["uq_target_repo (target_type, target_id, github_repo_id)", "idx_target", "idx_added_by"],
        "fks": [("added_by_user_id", "users", "id", "CASCADE")],
    },
    {
        "name": "communities",
        "group": "Community Hub",
        "columns": [
            ("id", "INT", "PK"),
            ("slug", "VARCHAR(120)", "UNIQUE, NOT NULL"),
            ("name", "VARCHAR(160)", "NOT NULL"),
            ("type", "ENUM", "city | university | dev_club | project_team | study_group | company | online"),
            ("description", "TEXT", ""),
            ("location", "VARCHAR(120)", "NULL"),
            ("website", "VARCHAR(255)", "NULL"),
            ("banner_url", "VARCHAR(500)", "NULL"),
            ("avatar_url", "VARCHAR(500)", "NULL"),
            ("owner_id", "INT", "NOT NULL"),
            ("member_count", "INT", "DEFAULT 0"),
            ("post_count", "INT", "DEFAULT 0"),
            ("is_public", "TINYINT(1)", "DEFAULT 1"),
            ("tags", "JSON", "NULL"),
            ("created_at", "DATETIME", ""),
            ("updated_at", "DATETIME", ""),
        ],
        "indexes": ["idx_type", "idx_location"],
        "fks": [("owner_id", "users", "id", "CASCADE")],
    },
    {
        "name": "community_memberships",
        "group": "Community Hub",
        "columns": [
            ("id", "INT", "PK"),
            ("community_id", "INT", "NOT NULL"),
            ("user_id", "INT", "NOT NULL"),
            ("role", "ENUM", "owner | admin | member"),
            ("joined_at", "DATETIME", ""),
        ],
        "indexes": ["uniq_member (community_id, user_id)"],
        "fks": [
            ("community_id", "communities", "id", "CASCADE"),
            ("user_id", "users", "id", "CASCADE"),
        ],
    },
    {
        "name": "community_posts",
        "group": "Community Hub",
        "desc": "Пости спільноти; linked_content_* — логічне посилання на hub без FK.",
        "columns": [
            ("id", "INT", "PK"),
            ("community_id", "INT", "NOT NULL"),
            ("author_id", "INT", "NOT NULL"),
            ("type", "ENUM", "discussion | pet_project | code_review | mentor_request | roadmap_request | team_search | event | announcement"),
            ("title", "VARCHAR(255)", "NOT NULL"),
            ("body", "TEXT", "NOT NULL"),
            ("metadata", "JSON", "NULL"),
            ("stack", "JSON", "NULL"),
            ("votes", "INT", "DEFAULT 0"),
            ("views", "INT", "DEFAULT 0"),
            ("comment_count", "INT", "DEFAULT 0"),
            ("status", "ENUM", "open | closed | filled"),
            ("linked_content_type", "VARCHAR(32)", "NULL, без FK"),
            ("linked_content_id", "INT", "NULL, без FK"),
            ("created_at", "DATETIME", ""),
            ("updated_at", "DATETIME", ""),
        ],
        "indexes": ["idx_community_type", "idx_linked_hub", "idx_author", "FULLTEXT(title, body)"],
        "fks": [
            ("community_id", "communities", "id", "CASCADE"),
            ("author_id", "users", "id", "CASCADE"),
        ],
    },
    {
        "name": "community_post_comments",
        "group": "Community Hub",
        "columns": [
            ("id", "INT", "PK"),
            ("post_id", "INT", "NOT NULL"),
            ("author_id", "INT", "NOT NULL"),
            ("parent_id", "INT", "NULL → self"),
            ("body", "TEXT", "NOT NULL"),
            ("votes", "INT", "DEFAULT 0"),
            ("created_at", "DATETIME", ""),
            ("updated_at", "DATETIME", ""),
        ],
        "indexes": ["idx_post"],
        "fks": [
            ("post_id", "community_posts", "id", "CASCADE"),
            ("author_id", "users", "id", "CASCADE"),
            ("parent_id", "community_post_comments", "id", "CASCADE"),
        ],
    },
    {
        "name": "mentor_profiles",
        "group": "Community Hub",
        "columns": [
            ("id", "INT", "PK"),
            ("user_id", "INT", "UNIQUE, NOT NULL"),
            ("is_active", "TINYINT(1)", "DEFAULT 1"),
            ("bio", "TEXT", ""),
            ("stack", "JSON", "NULL"),
            ("topics", "JSON", "NULL"),
            ("languages", "JSON", "NULL"),
            ("availability_hours_week", "INT", "DEFAULT 0"),
            ("price_note", "VARCHAR(160)", "NULL"),
            ("contact_method", "VARCHAR(160)", "NULL"),
            ("created_at", "DATETIME", ""),
            ("updated_at", "DATETIME", ""),
        ],
        "indexes": [],
        "fks": [("user_id", "users", "id", "CASCADE")],
    },
    {
        "name": "news_posts",
        "group": "Новини",
        "columns": [
            ("id", "INT", "PK"),
            ("title", "VARCHAR(255)", "NOT NULL"),
            ("summary", "VARCHAR(500)", "NOT NULL"),
            ("body", "TEXT", "NOT NULL"),
            ("slug", "VARCHAR(280)", "UNIQUE, NULL"),
            ("author_id", "INT", "NOT NULL"),
            ("category", "VARCHAR(32)", "DEFAULT tech"),
            ("published_at", "DATETIME", "NOT NULL"),
            ("is_pinned", "TINYINT(1)", "DEFAULT 0"),
            ("tags", "JSON", "NOT NULL"),
            ("views", "INT", "DEFAULT 0"),
            ("created_at", "DATETIME", "NOT NULL"),
            ("updated_at", "DATETIME", "NOT NULL"),
        ],
        "indexes": ["uq_news_slug", "idx_author", "idx_published", "idx_category", "FULLTEXT"],
        "fks": [("author_id", "users", "id", "CASCADE")],
    },
    {
        "name": "news_post_views",
        "group": "Новини",
        "columns": [
            ("id", "INT", "PK"),
            ("news_post_id", "INT", "NOT NULL"),
            ("viewer_key", "VARCHAR(96)", "NOT NULL"),
            ("viewed_at", "DATETIME", "NOT NULL"),
        ],
        "indexes": ["uq_news_viewer", "idx_news_post"],
        "fks": [("news_post_id", "news_posts", "id", "CASCADE")],
    },
    {
        "name": "news_comments",
        "group": "Новини",
        "columns": [
            ("id", "INT", "PK"),
            ("news_post_id", "INT", "NOT NULL"),
            ("author_id", "INT", "NOT NULL"),
            ("parent_id", "INT", "NULL → self"),
            ("body", "TEXT", "NOT NULL"),
            ("created_at", "DATETIME", ""),
            ("updated_at", "DATETIME", ""),
        ],
        "indexes": ["idx_news_comments_post"],
        "fks": [
            ("news_post_id", "news_posts", "id", "CASCADE"),
            ("author_id", "users", "id", "CASCADE"),
            ("parent_id", "news_comments", "id", "CASCADE"),
        ],
    },
    {
        "name": "news_polls",
        "group": "Новини",
        "desc": "Глобальні опитування (не прив’язані до одного поста).",
        "columns": [
            ("id", "INT", "PK"),
            ("slug", "VARCHAR(64)", "UNIQUE, NOT NULL"),
            ("title", "VARCHAR(255)", "NOT NULL"),
            ("description", "VARCHAR(500)", "NULL"),
            ("options", "JSON", "NOT NULL"),
            ("is_active", "TINYINT(1)", "DEFAULT 1"),
            ("created_at", "DATETIME", ""),
        ],
        "indexes": ["uq_news_poll_slug"],
        "fks": [],
    },
    {
        "name": "news_poll_votes",
        "group": "Новини",
        "columns": [
            ("id", "INT", "PK"),
            ("poll_id", "INT", "NOT NULL"),
            ("option_id", "VARCHAR(32)", "NOT NULL"),
            ("voter_key", "VARCHAR(96)", "NOT NULL"),
            ("user_id", "INT", "NULL"),
            ("voted_at", "DATETIME", ""),
        ],
        "indexes": ["uq_poll_voter (poll_id, voter_key)"],
        "fks": [
            ("poll_id", "news_polls", "id", "CASCADE"),
            ("user_id", "users", "id", "SET NULL"),
        ],
    },
]

RELATIONSHIPS = [
    ("users", "1:N", "user_repositories", "user_id"),
    ("users", "1:N", "questions", "author_id"),
    ("users", "1:N", "answers", "author_id"),
    ("questions", "1:N", "answers", "question_id"),
    ("users", "1:N", "votes", "user_id"),
    ("users", "1:N", "notifications", "user_id"),
    ("users", "1:N", "bookmarks", "user_id"),
    ("questions", "1:N", "bookmarks", "question_id"),
    ("questions", "1:N", "question_views", "question_id"),
    ("users", "1:N", "articles, guides, snippets, roadmaps, best_practices, faqs", "author_id"),
    ("articles", "1:N", "article_views", "article_id"),
    ("guides", "1:N", "guide_views", "guide_id"),
    ("snippets", "1:N", "snippet_views", "snippet_id"),
    ("roadmaps", "1:N", "roadmap_views", "roadmap_id"),
    ("best_practices", "1:N", "best_practice_views", "best_practice_id"),
    ("faqs", "1:N", "faq_views", "faq_id"),
    ("users", "1:N", "content_items", "author_id"),
    ("content_items", "1:N", "content_answers", "content_id"),
    ("content_items", "1:N", "content_views", "content_id"),
    ("content_items", "1:N", "content_bookmarks", "content_id"),
    ("users", "1:N", "content_bookmarks", "user_id"),
    ("users", "1:N", "content_linked_repos", "added_by_user_id"),
    ("users", "1:N", "communities", "owner_id"),
    ("communities", "1:N", "community_memberships", "community_id"),
    ("users", "1:N", "community_memberships", "user_id"),
    ("communities", "1:N", "community_posts", "community_id"),
    ("community_posts", "1:N", "community_post_comments", "post_id"),
    ("community_post_comments", "1:N", "community_post_comments", "parent_id"),
    ("users", "1:1", "mentor_profiles", "user_id"),
    ("users", "1:N", "news_posts", "author_id"),
    ("news_posts", "1:N", "news_post_views", "news_post_id"),
    ("news_posts", "1:N", "news_comments", "news_post_id"),
    ("news_comments", "1:N", "news_comments", "parent_id"),
    ("news_polls", "1:N", "news_poll_votes", "poll_id"),
]

MIGRATION_ORDER = [
    "users",
    "user_repositories",
    "questions",
    "answers",
    "votes",
    "notifications",
    "bookmarks",
    "question_views",
    "articles",
    "article_views",
    "guides",
    "guide_views",
    "snippets",
    "snippet_views",
    "roadmaps",
    "roadmap_views",
    "best_practices",
    "best_practice_views",
    "faqs",
    "faq_views",
    "content_items",
    "content_answers",
    "content_views",
    "content_bookmarks",
    "content_linked_repos",
    "communities",
    "community_memberships",
    "community_posts",
    "community_post_comments",
    "mentor_profiles",
    "news_posts",
    "news_post_views",
    "news_comments",
    "news_polls",
    "news_poll_votes",
]

POLYMORPHIC = [
    ("votes", "entity_type + entity_id → question | answer | content_items | content_answers"),
    ("notifications", "entity_type + entity_id (без FK)"),
    ("content_linked_repos", "target_type + target_id → legacy/hub таблиці"),
    ("community_posts", "linked_content_type + linked_content_id → hub (без FK)"),
]


def h(doc, text, level=1):
    return doc.add_heading(text, level=level)


def p(doc, text):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.size = Pt(11)
    run.font.name = "Calibri"


def bullet(doc, text):
    para = doc.add_paragraph(text, style="List Bullet")
    for run in para.runs:
        run.font.size = Pt(11)
        run.font.name = "Calibri"


def table_grid(doc, headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    for i, name in enumerate(headers):
        t.rows[0].cells[i].text = name
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            t.rows[ri + 1].cells[ci].text = str(val)
    doc.add_paragraph()


def write_table_section(doc, tbl):
    h(doc, tbl["name"], 2)
    if tbl.get("desc"):
        p(doc, tbl["desc"])
    rows = [[c[0], c[1], c[2]] for c in tbl["columns"]]
    table_grid(doc, ["Колонка", "Тип", "Обмеження / примітка"], rows)
    if tbl["indexes"]:
        p(doc, "Індекси: " + "; ".join(tbl["indexes"]))
    if tbl["fks"]:
        fk_rows = [
            [fk[0], f"{fk[1]}({fk[2]})", fk[3]]
            for fk in tbl["fks"]
        ]
        table_grid(doc, ["Колонка FK", "Посилання", "ON DELETE"], fk_rows)
    else:
        p(doc, "Зовнішні ключі: немає.")


def main():
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Cm(2)
    sec.bottom_margin = Cm(2)
    sec.left_margin = Cm(2.5)
    sec.right_margin = Cm(2)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("DevFlow — схема бази даних MySQL\n")
    r.bold = True
    r.font.size = Pt(22)
    r2 = title.add_run(f"Для регенерації БД іншим ШІ · {date.today().strftime('%d.%m.%Y')}\n")
    r2.font.size = Pt(12)
    r3 = title.add_run(f"Всього таблиць: {len(TABLES)}")
    r3.font.size = Pt(11)
    doc.add_paragraph()

    h(doc, "1. Загальні параметри", 1)
    bullet(doc, "СУБД: MySQL 8+, InnoDB")
    bullet(doc, "Кодування: utf8mb4, COLLATE utf8mb4_unicode_ci")
    bullet(doc, "Джерело правди: packages/backend/src/scripts/migrate.js")
    bullet(doc, "Теги та метадані: JSON-колонки")
    bullet(doc, "Перегляди: viewer_key / voter_key — user:id або visitor UUID")

    h(doc, "2. Порядок створення таблиць (залежності)", 1)
    for i, name in enumerate(MIGRATION_ORDER, 1):
        bullet(doc, f"{i}. {name}")

    h(doc, "3. Зв’язки між таблицями (FK)", 1)
    rel_rows = [[a, card, b, col] for a, card, b, col in RELATIONSHIPS]
    table_grid(doc, ["Батько", "Зв’язок", "Дочірня таблиця", "FK-колонка"], rel_rows)

    h(doc, "4. Поліморфні зв’язки (без FK)", 1)
    for tbl, note in POLYMORPHIC:
        bullet(doc, f"{tbl}: {note}")

    h(doc, "5. Перелік таблиць за групами", 1)
    groups = {}
    for t in TABLES:
        groups.setdefault(t["group"], []).append(t["name"])
    for grp, names in groups.items():
        h(doc, grp, 2)
        bullet(doc, ", ".join(names))

    h(doc, "6. Детальна специфікація кожної таблиці", 1)
    current_group = None
    for tbl in TABLES:
        if tbl["group"] != current_group:
            current_group = tbl["group"]
            h(doc, current_group, 2)
        write_table_section(doc, tbl)

    h(doc, "7. Примітки для регенерації", 1)
    bullet(doc, "Legacy: questions/answers/bookmarks/question_views — паралельно з content_items.")
    bullet(doc, "Hub існує в двох шарах: окремі таблиці (articles…) і універсальне content_items.")
    bullet(doc, "votes.entity_type включає content та content_answer для нового hub.")
    bullet(doc, "notifications.type/entity_type після міграції — VARCHAR(64), не ENUM.")
    bullet(doc, "news_posts.category додається ensureColumn (DEFAULT 'tech').")
    bullet(doc, "users.role на старих БД може бути VARCHAR замість ENUM — сумісність через ensureColumn.")

    doc.save(OUT)
    print(f"Збережено: {OUT}")


if __name__ == "__main__":
    main()
