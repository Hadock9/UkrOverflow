#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Генерація DOCX з детальним описом платформи DevFlow (станом на 2026)."""

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Cm

OUT = Path(__file__).resolve().parent.parent / "DevFlow-opys-saitu.docx"


def h(doc, text, level=1):
    return doc.add_heading(text, level=level)


def p(doc, text, bold=False):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.size = Pt(11)
    run.font.name = "Calibri"
    if bold:
        run.bold = True
    return para


def bullet(doc, text):
    para = doc.add_paragraph(text, style="List Bullet")
    for run in para.runs:
        run.font.size = Pt(11)
        run.font.name = "Calibri"


def table_headers(doc, headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, name in enumerate(headers):
        hdr[i].text = name
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            t.rows[ri + 1].cells[ci].text = str(val)
    doc.add_paragraph()


def main():
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Cm(2)
    sec.bottom_margin = Cm(2)
    sec.left_margin = Cm(2.5)
    sec.right_margin = Cm(2)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("DevFlow\n")
    r.bold = True
    r.font.size = Pt(22)
    r2 = title.add_run("Детальний опис функціоналу веб-платформи\n")
    r2.font.size = Pt(14)
    r3 = title.add_run(f"Дата: {date.today().strftime('%d.%m.%Y')}")
    r3.font.size = Pt(11)
    doc.add_paragraph()

    h(doc, "1. Загальна характеристика", 1)
    p(doc, (
        "DevFlow — україномовна платформа для розробників, що поєднує "
        "модель Q&A (Stack Overflow), knowledge hub (статті, гайди, сніпети, маршрути навчання), "
        "тематичні спільноти, каталог менторів, стрічку IT-новин у стилі DOU та інтеграцію "
        "штучного інтелекту (Google Gemini). Продакшн-інстанс: https://devflow.info"
    ))
    p(doc, (
        "Візуальний стиль — brutalism: чорно-білі контрасти, жовті/блакитні/зелені акценти, "
        "моноширинні підписи, товсті рамки та тіні box-shadow. Frontend: React + Vite; "
        "Backend: Node.js + Express; БД: MySQL; деплой: Docker (nginx + API + frontend build)."
    ))

    h(doc, "2. Навігація та структура сайту", 1)
    p(doc, "Головне меню (Header) містить:")
    bullet(doc, "ГОЛОВНА (/) — лендінг з оглядом розділів")
    bullet(doc, "ХАБ (/hub) — єдина стрічка knowledge hub")
    bullet(doc, "НОВИНИ (/news) — стрічка новин")
    bullet(doc, "ТЕГИ (/tags) — глобальний каталог тегів")
    bullet(doc, "КОМ'ЮНІТІ, МЕНТОРИ, РОЗРОБНИКИ (/devs)")
    bullet(doc, "Пошук у хедері — live-підказки під час введення (з 2 символів)")
    bullet(doc, "Увійти / Реєстрація або профіль + сповіщення (для авторизованих)")

    h(doc, "3. Сторінки та маршрути", 1)

    routes = [
        ("/", "Головна (MainPage)", "Hero, картки розділів, превʼю новин, блок «Свіже в хабі», бічна панель статистики"),
        ("/hub", "Хаб знань", "Уніфікована стрічка /api/content, фільтри за типом, сортування, пагінація, створення контенту"),
        ("/questions, /articles, …", "Фільтровані стрічки", "Той самий Home.jsx з автофільтром типу за URL"),
        ("/tags/:tag", "Контент за тегом", "Стрічка хабу з фільтром tag"),
        ("/tags", "Каталог тегів", "Хмара тегів, пошук, фільтри джерел, агрегація з усіх таблиць"),
        ("/search", "Глобальний пошук", "Live-підказки + повні результати FULLTEXT по хабу та спільнотах"),
        ("/news", "Стрічка новин", "Категорії (salary, career, tech…), пошук, sidebar: тренди, опитування, пульс ринку"),
        ("/news/:id", "Деталь новини", "Markdown-тіло, перегляди, коментарі, категорія"),
        ("/communities", "Спільноти", "Список міст/університетів/онлайн-груп"),
        ("/communities/:slug", "Сторінка спільноти", "Пости, учасники, membership"),
        ("/mentors", "Ментори", "Каталог профілів менторів"),
        ("/devs", "Розробники", "Каталог користувачів-профілів"),
        ("/users/:id, /profile", "Профіль", "Контент автора, GitHub-дані, вкладки хабу"),
        ("/login, /register", "Авторизація", "Email/пароль + GitHub OAuth"),
        ("/notifications", "Сповіщення", "Події: відповіді, голоси, коментарі в спільнотах"),
    ]
    table_headers(doc, ["URL", "Сторінка", "Опис"], routes)

    h(doc, "4. Knowledge Hub — типи контенту", 1)
    types = [
        ("Питання (question)", "Q&A з відповідями, голосування, accepted answer, перегляди"),
        ("Стаття (article)", "Довгі матеріали, excerpt, Markdown"),
        ("Гайд (guide)", "Міні-гайди, difficulty, estimated minutes"),
        ("Сніпет (snippet)", "Блоки коду + language"),
        ("Маршрут (roadmap)", "Кроки навчання (steps JSON), difficulty, тижні"),
        ("Best practice", "Правила, anti-patterns, category"),
        ("ЧаП (faq)", "Q&A пари (qa_pairs JSON)"),
        ("Пост спільноти", "У фіді /api/content як community_post"),
    ]
    table_headers(doc, ["Тип", "Можливості"], types)
    p(doc, "Єдиний API GET /api/content з параметрами contentType, sortBy, tag, authorId, search, page, limit.")

    h(doc, "5. Питання та відповіді", 1)
    bullet(doc, "Створення / редагування / видалення питань (автор або admin)")
    bullet(doc, "Відповіді з Markdown, голосування up/down")
    bullet(doc, "Прийнята відповідь (accepted)")
    bullet(doc, "Пов’язані пости спільнот за спільними тегами")
    bullet(doc, "Панель Linked GitHub Repos на сторінці питання")
    bullet(doc, "AI: підказка відповіді, TL;DR питання, схожі питання, резюме відповідей, рекомендації матеріалів")

    h(doc, "6. Спільноти (Communities)", 1)
    bullet(doc, "Типи: city, university, online та ін.")
    bullet(doc, "Slug-URL, опис, теги, website, location")
    bullet(doc, "Membership: owner, admin, member")
    bullet(doc, "Пости: pet_project, code_review, mentor_search, discussion тощо")
    bullet(doc, "Коментарі до постів, лічильники views/votes")
    bullet(doc, "Створення спільноти та постів авторизованими користувачами")

    h(doc, "7. Ментори", 1)
    bullet(doc, "Профіль ментора: стек, теми, формат сесій, контакт")
    bullet(doc, "Редагування власного профілю (/mentors/edit)")
    bullet(doc, "Перегляд каталогу /mentors")

    h(doc, "8. Стрічка новин", 1)
    p(doc, "Окремий модуль news_posts (не змішується з хабом):")
    bullet(doc, "Поля: title, summary, body, slug, category, tags, views, is_pinned, published_at")
    bullet(doc, "Категорії: salary, career, tech, community, events, ai")
    bullet(doc, "До ~400 демо-новин (seed:news:large), тематика українського IT / DOU")
    bullet(doc, "Sidebar: «Популярне за тиждень», «Зарплати», топ тегів, опитування грейду")
    bullet(doc, "Коментарі під новинами (news_comments)")
    bullet(doc, "Створення новин: admin / moderator (/news/new)")

    h(doc, "9. Пошук", 1)
    bullet(doc, "GET /api/search/live — миттєві підказки (хаб + новини + теги)")
    bullet(doc, "GET /api/search/global — повнотекстовий пошук по всіх типах хабу + community_post")
    bullet(doc, "GET /api/search — legacy пошук питань/відповідей")
    bullet(doc, "UI: dropdown у хедері, сторінка /search з автооновленням URL")

    h(doc, "10. Теги", 1)
    bullet(doc, "GET /api/tags — агрегація тегів з questions, articles, guides, snippets, roadmaps, faqs, news, communities")
    bullet(doc, "Сторінка /tags: хмара, фільтри, сортування, посилання на /tags/:name та /news?tag=")

    h(doc, "11. AI-інтеграція (Google Gemini)", 1)
    ai_rows = [
        ("POST /api/ai/suggest-answer", "Чернетка відповіді на питання (JWT)"),
        ("POST /api/ai/suggest-tags", "Автопідбір тегів при створенні"),
        ("POST /api/ai/summarize", "TL;DR довгого питання"),
        ("POST /api/ai/summarize-answers", "Резюме дискусії"),
        ("GET /api/ai/similar-questions/:id", "Схожі питання"),
        ("GET /api/ai/related-content/:id", "Рекомендації статей/гайдів/roadmap"),
        ("POST /api/ai/check-duplicate", "Перевірка дублікатів перед публікацією"),
        ("POST /api/ai/analyze-question", "Оцінка якості формулювання"),
        ("POST /api/ai/generate-roadmap", "Генерація roadmap за стеком (JWT)"),
        ("POST /api/ai/moderate", "Модерація SPAM/TOXIC/LOW_QUALITY"),
    ]
    table_headers(doc, ["Endpoint", "Призначення"], ai_rows)
    p(doc, "UI: AIQuestionCoach (NewQuestion), AIAnswersSummary, AIRelatedPosts, AISimilarQuestions, AIQuestionSummary, AIAssistant, AITagSuggester, AIRoadmapGenerator.")

    h(doc, "12. Авторизація та користувачі", 1)
    bullet(doc, "Реєстрація email + пароль (bcrypt)")
    bullet(doc, "GitHub OAuth: login, link account, callback /auth/callback")
    bullet(doc, "Ролі: user, moderator, admin")
    bullet(doc, "Репутація (reputation), аватар, bio, location, website")
    bullet(doc, "Синхронізація GitHub: профіль, stack, contributions, badges, repos")
    bullet(doc, "Каталог /devs та адмін-список /users")

    h(doc, "13. Сповіщення та закладки", 1)
    bullet(doc, "Типи: question_answer, answer_comment, vote, mention, community_post_comment тощо")
    bullet(doc, "NotificationBell у хедері, сторінка /notifications")
    bullet(doc, "Закладки на питання (bookmarks API)")

    h(doc, "14. Статистика", 1)
    bullet(doc, "GET /api/stats/overview — загальні лічильники")
    bullet(doc, "GET /api/stats/top-tags, top-users, recent-activity, unanswered")
    bullet(doc, "StatsSidebar на головній: питання, відповіді, користувачі, топ теги/користувачі")

    h(doc, "15. Технічний стек", 1)
    stack = [
        ("Frontend", "React 18, Vite, React Router, TanStack Query, Axios, marked+DOMPurify, brutalism CSS"),
        ("Backend", "Node.js ESM, Express, express-validator, JWT, mysql2, WebSocket (ws)"),
        ("Shared", "Пакет shared — типи контенту, шляхи"),
        ("Mediator", "Патерн Mediator для подій UI (окремий пакет)"),
        ("AI", "Gemini REST API (Flash/Pro), GEMINI_API_KEY"),
        ("Deploy", "docker-compose: nginx (SSL Let's Encrypt), api, web, MySQL"),
    ]
    table_headers(doc, ["Шар", "Технології"], stack)

    h(doc, "16. База даних (основні таблиці)", 1)
    tables = [
        "users, questions, answers, votes, bookmarks, notifications",
        "articles, guides, snippets, roadmaps, best_practices, faqs, content_items (hub)",
        "communities, community_memberships, community_posts, community_post_comments",
        "mentor_profiles, news_posts, news_comments, news_polls, news_poll_votes, news_post_views",
        "user_repositories, content_linked_repos, github webhooks",
    ]
    for t in tables:
        bullet(doc, t)

    h(doc, "17. Скрипти наповнення даними", 1)
    seeds = [
        ("npm run seed", "~10 users, 490 hub units, 700 answers"),
        ("npm run seed:communities:large", "15 спільнот, багато постів/членів"),
        ("npm run seed:news:large", "~400 новин"),
        ("npm run seed:news-poll", "Опитування грейду (DOU-style)"),
        ("npm run seed:user -- --user N", "Демо-контент для користувача N"),
        ("npm run seed:notifications", "Демо-сповіщення"),
        ("npm run migrate", "Створення/оновлення схеми БД"),
    ]
    table_headers(doc, ["Команда", "Що робить"], seeds)

    h(doc, "18. Обмеження та відомі нюанси", 1)
    bullet(doc, "AI вимагає GEMINI_API_KEY; без ключа backend може не стартувати")
    bullet(doc, "Rate limit API вимкнено за замовчуванням (увімкнення через API_RATE_LIMIT_ENABLED=1)")
    bullet(doc, "Mediator Visualizer вимкнено на проді (був debug UI)")
    bullet(doc, "На VPS потрібні узгоджені FRONTEND_URL та GitHub OAuth callback")
    bullet(doc, "VITE_API_URL порожній на проді — запити йдуть через той самий origin (/api)")

    h(doc, "19. Висновок", 1)
    p(doc, (
        "На момент опису DevFlow — повноцінна дипломна платформа з knowledge hub, спільнотами, "
        "новинами, менторством, глобальним пошуком, каталогом тегів та розгалуженою AI-інтеграцією. "
        "Сайт орієнтований на українську IT-спільноту та демонструє сучасний full-stack підхід "
        "з Docker-деплоєм і production-ready доменом devflow.info."
    ))

    doc.save(str(OUT))
    print(f"Збережено: {OUT}")


if __name__ == "__main__":
    main()
